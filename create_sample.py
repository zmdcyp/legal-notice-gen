"""创建示例模板和示例 Excel，方便测试。"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl


def create_sample_template():
    doc = Document()

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("律  师  函")
    run.bold = True
    run.font.size = Pt(22)

    doc.add_paragraph()

    # 正文
    doc.add_paragraph(f"致：{{{{收件人}}}}")
    doc.add_paragraph(f"地址：{{{{收件地址}}}}")
    doc.add_paragraph()

    body = doc.add_paragraph()
    body.add_run(
        "本律师受{{委托人}}的委托，就{{案由}}事宜，"
        "依法向贵方发出本律师函。"
    )

    doc.add_paragraph()

    body2 = doc.add_paragraph()
    body2.add_run(
        "经查，贵方于{{事实日期}}{{事实描述}}。"
        "上述行为已构成对我方委托人合法权益的侵害。"
    )

    doc.add_paragraph()

    body3 = doc.add_paragraph()
    body3.add_run(
        "现我方委托人要求贵方在收到本函之日起{{期限}}内，{{要求内容}}。"
        "逾期未予回复或处理，我方委托人将依法采取进一步法律行动，届时由此产生的一切法律后果由贵方承担。"
    )

    doc.add_paragraph()
    doc.add_paragraph("特此函告。")
    doc.add_paragraph()

    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign.add_run("{{律所名称}}")

    # 印章占位符
    stamp_p = doc.add_paragraph()
    stamp_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    stamp_p.add_run("{{stamp}}")

    sign2 = doc.add_paragraph()
    sign2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign2.add_run("律师：{{律师姓名}}")

    # 签名占位符
    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign_p.add_run("{{sign}}")

    sign3 = doc.add_paragraph()
    sign3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign3.add_run("日期：{{发函日期}}")

    doc.save("sample_template.docx")
    print("已创建: sample_template.docx")


def create_sample_excel():
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "律师函数据"

    headers = ["收件人", "收件地址", "委托人", "案由", "事实日期",
               "事实描述", "期限", "要求内容", "律所名称", "律师姓名", "发函日期"]
    ws.append(headers)

    ws.append([
        "张三", "北京市朝阳区XX路XX号",
        "李四", "合同纠纷", "2025年6月15日",
        "未按合同约定支付货款人民币50万元整",
        "15个工作日", "支付拖欠货款人民币50万元整及违约金",
        "XX律师事务所", "王律师", "2026年4月1日"
    ])

    ws.append([
        "某某公司", "上海市浦东新区XX大道XX号",
        "赵六", "侵权纠纷", "2025年8月20日",
        "未经授权使用我方委托人的注册商标",
        "10个工作日", "立即停止侵权行为并赔偿经济损失",
        "XX律师事务所", "王律师", "2026年4月1日"
    ])

    wb.save("sample_data.xlsx")
    print("已创建: sample_data.xlsx")


if __name__ == "__main__":
    create_sample_template()
    create_sample_excel()
